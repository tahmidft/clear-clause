#!/usr/bin/env python3
"""Build sample contract DOCX and TXT files for ClearClause demos."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from contract_content import LEGACY_ALIASES, VARIANTS

SAMPLES_DIR = Path(__file__).resolve().parent


def sections_to_plain_text(sections: list[tuple[str, str]]) -> str:
    parts: list[str] = []
    for heading, body in sections:
        parts.append(heading)
        parts.append(body)
        parts.append("")
    return "\n".join(parts).strip() + "\n"


def write_docx(sections: list[tuple[str, str]], out_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run(sections[0][0])
    run.bold = True
    run.font.size = Pt(14)

    if "\n" in sections[0][1]:
        first_lines = sections[0][1].split("\n", 1)
        if len(first_lines) > 1 and first_lines[0].startswith("("):
            sub = doc.add_paragraph()
            sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            sub.add_run(first_lines[0]).italic = True

    doc.add_paragraph()

    for heading, body in sections[1:]:
        h = doc.add_paragraph()
        h.add_run(heading).bold = True
        doc.add_paragraph(body)
        doc.add_paragraph()

    doc.save(out_path)


def build_variant(name: str, sections: list[tuple[str, str]]) -> None:
    stem = f"{name}-freelance-contract-sample"
    txt_path = SAMPLES_DIR / f"{stem}.txt"
    docx_path = SAMPLES_DIR / f"{stem}.docx"
    txt_path.write_text(sections_to_plain_text(sections), encoding="utf-8")
    write_docx(sections, docx_path)
    print(f"Wrote {txt_path.name} ({len(txt_path.read_text())} chars)")
    print(f"Wrote {docx_path.name}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Build ClearClause sample contracts")
    parser.add_argument(
        "variants",
        nargs="*",
        choices=[*VARIANTS.keys(), "all", *LEGACY_ALIASES.keys()],
        default=["all"],
        help="Which samples to build (default: all)",
    )
    args = parser.parse_args()
    requested = args.variants or ["all"]
    if "all" in requested:
        names = list(VARIANTS.keys())
    else:
        names = [LEGACY_ALIASES.get(v, v) for v in requested]

    for name in names:
        if name not in VARIANTS:
            raise SystemExit(f"Unknown variant: {name}")
        build_variant(name, VARIANTS[name])

    # Backward-compatible legacy filenames (same as bad)
    if "bad" in names or "all" in requested:
        bad = VARIANTS["bad"]
        legacy_txt = SAMPLES_DIR / "freelance-design-contract-sample.txt"
        legacy_docx = SAMPLES_DIR / "freelance-design-contract-sample.docx"
        legacy_txt.write_text(sections_to_plain_text(bad), encoding="utf-8")
        write_docx(bad, legacy_docx)
        print(f"Wrote legacy {legacy_txt.name}")
        print(f"Wrote legacy {legacy_docx.name}")


if __name__ == "__main__":
    main()
